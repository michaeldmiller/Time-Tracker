# Welcome to Time Tracker, version 0.8.3

# changelog:
# slight alteration to career coloring

# import necessary libraries


import matplotlib.pyplot as pyplot
import matplotlib.patches as patches
import numpy as np
import time
import calendar
import colorsys
import copy

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.section import WD_SECTION
from docx.shared import Inches
from docx.shared import Pt

from openpyxl import Workbook
from openpyxl import load_workbook


def color_category_sorter(category):
    """Takes activity category and generates its corresponding color
        str->tuple"""
    color_dict = {"sleep": [0, 29, 105], "entertainment": [255, 48, 48], "career": [10, 75, 158],
                  "outdoor": [255, 220, 90], "project": [15, 251, 251], "socialization": [255, 99, 196],
                  "grooming": [166, 255, 123], "culinary": [255, 191, 30], "relaxation": [91, 173, 68],
                  "fitness": [255, 251, 67], "housework": [191, 131, 0], "maintenance": [118, 55, 0],
                  "volunteering": [140, 20, 86], "transit": [218, 218, 218], "other": [80, 180, 155],
                  "missing": [87, 87, 87]}
    intermediate_color = color_dict[category]
    intermediate_color_counter = 0
    while intermediate_color_counter < 3:
        if intermediate_color[intermediate_color_counter] == 0:
            pass
        else:
            intermediate_color[intermediate_color_counter] = intermediate_color[intermediate_color_counter] / 255
        intermediate_color_counter += 1
    return intermediate_color


def inner_color_creator(rgb_output, lighten_value=0.7):
    """Takes color_category_sorter RGB Tuple, adds lightened alpha channel
        tup->tup"""
    start_r = rgb_output[0]
    start_g = rgb_output[1]
    start_b = rgb_output[2]
    start_hsv = list(colorsys.rgb_to_hsv(start_r, start_g, start_b))
    start_hsv[1] = start_hsv[1] * lighten_value
    new_h = start_hsv[0]
    new_s = start_hsv[1]
    new_v = start_hsv[2]
    new_rgb_output = colorsys.hsv_to_rgb(new_h, new_s, new_v)
    return new_rgb_output


def hour_and_minute_converter(minutes):
    """takes a number of minutes and converts it into hours and minutes
        int->str"""
    minutes = int(minutes)
    n_hours, n_minutes = minutes // 60, minutes % 60
    n_hours = str(n_hours)
    n_minutes = str(n_minutes)
    if len(n_minutes) < 2:
        n_minutes = "0" + n_minutes
    time_output = n_hours + ":" + n_minutes
    return time_output


def report_day_converter(report_day):
    """takes report day (str in format: YYYY/MM/DD) and converts to
        format: YYYY-MM-DD
        str->str"""
    new_report_day = report_day.replace("/", "-")
    return new_report_day


# Begin data acquisition section

# infinite loop with escape, get time range and category of each entry,
# append to list

# get date for chart
print("""
This program will generate charts representing time usage
on a given day. Each day can be divided into activity blocks,
with a start time, an end time, and an activity type (or category).

For what day is this report being generated? Please specify in YYYY/MM/DD""")

while True:
    report_day = str(input())
    try:
        report_day_intermediate = report_day.replace("/", " ")
        report_day_object = time.strptime(report_day_intermediate, "%Y %m %d")
        break
    except:
        print("I'm sorry, that input was not valid\n")
        continue

# explain category system
print("""
This program will now request that you input the details for
the activity blocks which best represent the day. For each activity
block, start time, end time, and activity type will be requested.
Valid activity types are: (note - these are case insensitive)
'Sleep'
'Entertainment'
'Career' \t\t(anything relating to work, career, or job search)
'Outdoor' \t\t(anything related to work outdoors, e.g. yard work, construction)
'Project' \t\t(anything relating to a project/hobby of personal interest)
'Socialization'
'Grooming'
'Culinary'
'Relaxation' \t\t(w/o active entertainment, such as downtime, thinking, etc)
'Fitness'
'Housework'
'Maintenance'
'Volunteering'
'Transit'
'Other'

Starting with version 0.8.1, the entire day's block start and end times, as
well as categories, will be requested in a single block in the following format:
00:00 - 01:00 = Relaxation
Where the first 5 characters are the hh:mm of the start time,
the next 3 characters are space, dash, space,
the following block of 5 characters are the hh:mm of the end time,
the next 3 characters are space, equal sign, space, 
and the remaining characters are a valid category (case insensitive)

These blocks are to be placed in a text file named "day_summary.txt"
in the same directory as the main program.\n
""")
# initialize variables
entry_number = 1
activities_accum = []
valid_categories = ['sleep', 'entertainment', 'career', 'outdoor', 'project',
                    'socialization', 'grooming', 'culinary', 'relaxation', 'fitness',
                    'housework', 'maintenance',  'volunteering', 'transit', 'other', 'missing']

# create variable containing times (in minutes), used by an activity block


# print(unusedTimes)
fake_date = "2000 01 01 "
would_like_to_leave = False

# Check for file of all inputs
while True:
    print("""
Please ensure the existence of a valid text file named day_summary.txt
in the program's directory. Has this been done? (Yes/No)
    """)
    confirmation_of_file = str(input()).lower()
    block_input_error = False
    unused_times = []
    day_blocks_list = []
    i = 0
    while i < 1440:
        unused_times.append(i)
        i += 1
    if confirmation_of_file == "yes":
        try:
            day_blocks_file = open("day_summary.txt", "r")
            for line in day_blocks_file:
                day_blocks_list.append(line)
            print(day_blocks_list)
            day_blocks_file.close()

        except:
            print("Please ensure the appropriate file exists.")
            block_input_error = True
            continue

        for block in day_blocks_list:
            # print("block before", block)
            block = block.replace(" ", "")
            block = block.replace("=", "")
            block = block.replace("-", "")
            block = block.replace("\n", "")
            # print("block after", block)
            block_start_time = block[0:5]
            # print(block_start_time)
            if len(block_start_time) != 5:
                print("I'm sorry, that input was not valid\n")
                block_input_error = True
                break
            try:
                block_start_time = block_start_time.replace(":", " ")
                intermediate_start_time = fake_date + block_start_time
                block_start_time_object = time.strptime(intermediate_start_time, "%Y %m %d %H %M")
                block_start_time = block_start_time.replace(" ", "")
            except:
                try:
                    block_start_time = block_start_time.replace(";", " ")
                    intermediate_start_time = fake_date + block_start_time
                    block_start_time_object = time.strptime(intermediate_start_time, "%Y %m %d %H %M")
                    block_start_time = block_start_time.replace(" ", "")
                except:
                    print("I'm sorry, that input was not valid \n")
                    block_input_error = True
                    break

            # ensure that time was not 11:59, as valid block cannot be created if it is

            # convert times to minutes of a day (i.e. out of 1440)
            start_time_minutes = ((int(block_start_time[0:2]) * 60) + int(block_start_time[2:4]))
            if start_time_minutes == 1439:
                print("An activity cannot begin at 23:59 in this calculator, please enter a new time")
                block_input_error = True
                break

            # ensure time is not in a block that has been used, and that a valid solution exists
            if start_time_minutes in unused_times and start_time_minutes + 1 in unused_times:
                pass
            else:
                print("An activity cannot begin during a time that has already been used, please enter a new time")
                block_input_error = True
                break

            # deal with end time
            block_end_time = block[5:10]
            # print(block_end_time)

            if len(block_end_time) != 5:
                print("I'm sorry, that input was not valid\n")
                block_input_error = True
                break
            try:
                block_end_time = block_end_time.replace(":", " ")
                intermediate_end_time = fake_date + block_end_time
                block_end_time_object = time.strptime(intermediate_end_time, "%Y %m %d %H %M")
                block_end_time = block_end_time.replace(" ", "")
            except:
                try:
                    block_end_time = block_end_time.replace(";", " ")
                    intermediate_end_time = fake_date + block_end_time
                    block_end_time_object = time.strptime(intermediate_end_time, "%Y %m %d %H %M")
                    block_end_time = block_end_time.replace(" ", "")
                except:
                    print("I'm sorry, that input was not valid\n")
                    block_input_error = True
                    break

            # ensure end does not occur before the beginning (time paradox)
            end_time_minutes = ((int(block_end_time[0:2]) * 60) + int(block_end_time[2:4]))
            if end_time_minutes > start_time_minutes:
                pass
            else:
                print("The end cannot occur before the beginning, please retry\n")
                block_input_error = True
                break

            # ensure time is not in a block that has been used,
            if end_time_minutes in unused_times:
                pass
            else:
                print("An activity cannot occur during a time that has already been used, please enter a new time")
                block_input_error = True
                break

            # ensure that interval does not supersede another block
            used_times_counter = 0
            used_times_list = []
            used_times_counter_target = end_time_minutes - start_time_minutes + 1
            possible_time_error_flag = False
            while used_times_counter < used_times_counter_target:
                used_times_list.append(used_times_counter + start_time_minutes)
                used_times_counter += 1
            for possible_time in used_times_list:
                if possible_time not in unused_times:
                    possible_time_error_flag = True

            if possible_time_error_flag == False:
                pass
            else:
                print("An activity block cannot supersede another activity block")
                block_input_error = True
                continue

            block_category = block[10:]
            # print(block_category)

            # prevent human idiocy
            block_category = block_category.lower()
            if block_category not in valid_categories:
                print("I'm sorry, that input was not valid\n")
                block_input_error = True
                break

            # with all categories acquired, store data
            # print(start_time_minutes)
            # print(end_time_minutes)
            # print(block_category)
            # create dictionary list
            block_summary_list = ([start_time_minutes, end_time_minutes, block_category])
            # append to dictionary
            activities_accum.append(block_summary_list)

            # remove block times from unused times
            for a_time in used_times_list:
                if a_time in unused_times:
                    unused_times.remove(a_time)

        if block_input_error == False:
            chronological_activities = sorted(activities_accum, key=lambda entry: entry[0])
            # print(chronological_activities)
            # print(unused_times)
            # print(report_day)
            break

    else:
        print("Please confirm to continue.")
        continue


# lines here creating data that would be generated or established by main time tracker program

##chronologicalActivities = [[30, 630, 'sleep'], [645, 675, 'grooming'], [676, 694, 'culinary'], [705, 735, 'socialization'],
##                           [750, 984, 'outdoor'], [990, 1010, 'grooming'], [1020, 1067, 'culinary'], [1080, 1115, 'relaxation'],
##                           [1116, 1143, 'relaxation'], [1150, 1336, 'entertainment'], [1345, 1360, 'career'],
##                           [1361, 1390, 'project'], [1395, 1435, 'fitness']]
##
##unusedTimes = [0, 1, 2, 3, 4, 5, 6, 7, 8, 9, 10, 11, 12, 13, 14, 15, 16, 17, 18, 19, 20, 21, 22, 23, 24, 25, 26, 27, 28, 29,
##               631, 632, 633, 634, 635, 636, 637, 638, 639, 640, 641, 642, 643, 644, 695, 696, 697, 698, 699, 700, 701, 702,
##               703, 704, 736, 737, 738, 739, 740, 741, 742, 743, 744, 745, 746, 747, 748, 749, 985, 986, 987, 988, 989, 1011,
##               1012, 1013, 1014, 1015, 1016, 1017, 1018, 1019, 1068, 1069, 1070, 1071, 1072, 1073, 1074, 1075, 1076, 1077, 1078,
##               1079, 1144, 1145, 1146, 1147, 1148, 1149, 1337, 1338, 1339, 1340, 1341, 1342, 1343, 1344, 1391, 1392, 1393, 1394,
##               1436, 1437, 1438, 1439]
##report_day = "2020/06/15"


# create first chart, the line chart
figure1, ax1 = pyplot.subplots()
ax1.plot()

# set the x axis to cover 0-1440
ax1.set_xlim(-25, 1465)
# set the y axis to cover 0-1
ax1.set_ylim(0, 1)

# set the aspect ratio
ax1.set_aspect(aspect=300)

# make the axes invisible
ax1.axis('off')

# add labels
hour_label_counter = 0

hour_label_x = 0
time_label_dict = {0: '12:00 AM', 1: "1:00 AM", 2: "2:00 AM", 3: "3:00 AM", 4: "4:00 AM",
                   5: '5:00 AM', 6: "6:00 AM", 7: "7:00 AM", 8: "8:00 AM", 9: "9:00 AM",
                   10: '10:00 AM', 11: "11:00 AM", 12: "12:00 PM", 13: "1:00 PM", 14: "2:00 PM",
                   15: '3:00 PM', 16: '4:00 PM', 17: '5:00 PM', 18: '6:00 PM', 19: '7:00 PM',
                   20: '8:00 PM', 21: '9:00 PM', 22: '10:00 PM', 23: '11:00 PM', 24: '12:00 AM'}

while hour_label_counter < 25:
    ax1.add_patch(patches.Rectangle((hour_label_x, 0.25), 1.2, 0.75, facecolor='k', edgecolor=None))
    pyplot.text(hour_label_x, 0.19, s=time_label_dict[hour_label_counter],
                horizontalalignment='center', fontsize=2.5)
    hour_label_x += 60
    hour_label_counter += 1

ax1.add_patch(patches.Rectangle((-25, 0.18), 1490, 0.0035, facecolor='k', edgecolor=None))
ax1.add_patch(patches.Rectangle((-25, 0.22), 1490, 0.0035, facecolor='k', edgecolor=None))

# create missing time activity blocks
place_in_unused_times = 0
try:
    current_block_start = unused_times[0]

    while place_in_unused_times < len(unused_times) - 1:
        # see if there is a gap, ie an activity block, by subtracting the value in
        # the next place from the current place
        if unused_times[place_in_unused_times + 1] - unused_times[place_in_unused_times] != 1:
            # create new missing activity block
            new_missing_activity_block = [current_block_start, int(unused_times[place_in_unused_times]), 'missing']

            # append new block to master list
            chronological_activities.append(new_missing_activity_block)
            # set new block starting point
            current_block_start = unused_times[place_in_unused_times + 1]

        place_in_unused_times += 1
        if place_in_unused_times == len(unused_times) - 1:
            new_missing_activity_block = [current_block_start, int(unused_times[place_in_unused_times]), 'missing']
            chronological_activities.append(new_missing_activity_block)
except:
    None

# re-sort chronologicalActivities
chronological_activities = sorted(chronological_activities, key=lambda entry: entry[0])

# add data rectangles
for final_activity_block in chronological_activities:
    block_category = final_activity_block[2]
    block_color = color_category_sorter(block_category)
    ax1.add_patch(
        patches.Rectangle((final_activity_block[0], 0.510), int(final_activity_block[1] - final_activity_block[0] + 1), 0.2,
                          facecolor=block_color, edgecolor=None))

# add date
pyplot.text(1465, 0.09, s="Report Date: %s" % report_day, horizontalalignment='right', fontsize=4)

# create legend
pyplot.text(-25, 0.07, s="Legend: ", horizontalalignment='left', fontsize=4)
category_list = ["sleep", "maintenance", "career", "project", "fitness",
                 "outdoor", "entertainment", "relaxation", "socialization",
                 "grooming", "culinary", "housework", "volunteering",
                 "other", "transit", "missing"]

legend_x_coord = 60
legend_y_coord = 0.10
for legend_proto_entry in category_list:
    legend_color = color_category_sorter(legend_proto_entry)
    legend_title = legend_proto_entry
    if legend_title == "outdoor":
        legend_title = "Outdoor Work"
    ax1.add_patch(patches.Rectangle((legend_x_coord, legend_y_coord), 15, 0.05,
                                    facecolor=legend_color, edgecolor=None))
    pyplot.text(legend_x_coord + 20, legend_y_coord + 0.005, s=legend_title.title(),
                horizontalalignment='left', fontsize=3.25)
    legend_y_coord -= 0.07
    # if reached bottom of column, start next column
    if round(legend_y_coord, 2) == -0.04:
        legend_x_coord += 150
        legend_y_coord = 0.10
# save plot 1
pyplot.savefig("day_report", dpi=800, orientation='landscape', bbox_inches='tight')

# create second plot
figure2, ax2 = pyplot.subplots()

wedge_width = 0.3

# create array containing the duration of all experience blocks
# print(chronologicalActivities)

categorized_activities = sorted(chronological_activities, key=lambda entry: entry[2])
# print(categorizedActivities)

# pull out category order
identified_activities = []
for entry in categorized_activities:
    if entry[2] not in identified_activities:
        identified_activities.append(entry[2])

# create Activities list divorced from start and end time, only duration
duration_activities = []
for block in categorized_activities:
    event_duration = block[1] - block[0] + 1
    duration_activities.append([event_duration, block[2]])

# create duration array
# separate durations by category

current_set = []
duration_proto_list = []
current_category = duration_activities[0][1]
duration_activities_loop_counter = 0

for pair in duration_activities:
    if pair[1] == current_category:
        current_set.append(pair[0])
    else:
        duration_proto_list.append(current_set)
        current_set = [pair[0]]

    current_category = pair[1]
    duration_activities_loop_counter += 1
    if duration_activities_loop_counter == len(duration_activities):
        duration_proto_list.append(current_set)

# next step in creating array: find the largest number of entries in
# a single category

max_num_repeats = 1
for listing in duration_proto_list:
    if len(listing) > max_num_repeats:
        max_num_repeats = len(listing)

# following step: add junk 0's to categories with less entries
# (this makes the matrix 'regular' permitting matrix algebra
# and ensures matplotlib won't flip out)

duration_array = duration_proto_list
for listing in duration_array:
    while True:
        if len(listing) < max_num_repeats:
            listing.append(0)
        else:
            break

# sort in descending order by highest duration category
current_locator_in_list = 0
for nextListing in duration_array:
    nextListing.append(identified_activities[current_locator_in_list])
    current_locator_in_list += 1
    totalDuration = sum(nextListing[0:len(nextListing) - 1])
    nextListing.append(totalDuration)

duration_array = sorted(duration_array, key=lambda entry: entry[-1], reverse=True)
# get durations for use in table
durations = []
for entry in duration_array:
    durations.append(entry[-1])

duration_array_counter = 0
identified_activities = []

for listing in duration_array:
    identified_activities.append(listing[-2])
    listing = listing[:-2]
    duration_array[duration_array_counter] = listing
    duration_array_counter += 1

# print(durationArray)
duration_array = np.array(duration_array)

# set outer ring's colors, creating custom cmap
outer_ring_colors = []
for category in identified_activities:
    outer_ring_colors.append(color_category_sorter(category))
outer_ring_colors = np.array(outer_ring_colors)

# set inner ring's colors, creating custom cmap
inner_ring_colors = []
for category2 in identified_activities:
    current_color = color_category_sorter(category2)
    temp_counter = 0
    lighten_val = 0.6
    while temp_counter < max_num_repeats:
        inner_ring_colors.append(inner_color_creator(current_color, lighten_val))
        if lighten_val >= 0.25:
            lighten_val -= 0.2
        temp_counter += 1

inner_ring_colors = np.array(inner_ring_colors)

# create the labels

pie_chart_labels = copy.deepcopy(identified_activities)
pie_chart_label_counter = 0
while pie_chart_label_counter < len(pie_chart_labels):
    pie_chart_labels[pie_chart_label_counter] = pie_chart_labels[pie_chart_label_counter].capitalize()
    pie_chart_label_counter += 1

explode = []
for item in pie_chart_labels:
    explode.append(0.2)
explode = tuple(explode)

# create pie chart outer ring
ax2.pie(duration_array.sum(axis=1), labels=None, radius=1,
        startangle=90, colors=outer_ring_colors, autopct=None,
        pctdistance=0.85, explode=None,
        wedgeprops=dict(width=wedge_width, edgecolor='w'))
# create pie chart inner ring
ax2.pie(duration_array.flatten(), radius=1 - wedge_width, colors=inner_ring_colors,
        startangle=90,
        wedgeprops=dict(width=wedge_width, edgecolor='w'))
# ax2.set_aspect("equal")
ax2.legend(loc='upper left',
           labels=['%s, %1.1f%%' % (l, s) for l, s in
                   zip(pie_chart_labels, (duration_array.sum(axis=1) / duration_array.sum() * 100))],
           bbox_to_anchor=(-0.15, 0.75), bbox_transform=pyplot.gcf().transFigure)

pyplot.savefig("pie_report", dpi=500, orientation='landscape', bbox_inches='tight')

# create values for a table of category durations

table_categories = copy.deepcopy(pie_chart_labels)
table_values = durations

# convert durations into hours and minutes
table_durations = []
for duration in durations:
    table_durations.append(hour_and_minute_converter(duration))

# begin saving process using python.docx module
# determine what day this is
# documents are Week1.docx, Week2.docx, Week3.docx,

clean_up_your_mess = False
somethings_wrong = False

week1_exists = False
week2_exists = False
week3_exists = False
week4_exists = False

# see if there is a Week1.docx,
# if there is, open it and get the number of sections
try:
    document1 = Document('Week1.docx')
    week1_exists = True
    num_sections_wk1 = len(document1.sections)
except:
    None

# try to open Week2.docx and get its number of sections
try:
    document2 = Document('Week2.docx')
    week2_exists = True
    num_sections_wk2 = len(document2.sections)
except:
    None

# try to open Week3.docx and get its number of sections
try:
    document3 = Document('Week3.docx')
    week3_exists = True
    num_sections_wk3 = len(document3.sections)
except:
    None

# try to open Week4.docx and get its number of sections
try:
    document4 = Document('Week4.docx')
    week4_exists = True
    num_sections_wk4 = len(document4.sections)
except:
    None

# massive logic sorter for the combination of present files
# program will only proceed if an appropriate file combination is present
# valid combinations are:
#   Nothing present
#   Week1 present without Week2, Week3, or Week4
#   Week1 and Week2 present without Week3 or Week4
#   Week1, Week2, and Week3 present without Week4
#   Week1, Week2, Week3, and Week4 present

# invalid combination triggers are:
#   Week2 present without Week1
#   Week3 present without Week1
#   Week3 present without Week2
#   Week4 present without Week1
#   Week4 present without Week2
#   Week4 present without Week3

# to ensure file validity the program will loop through all
# invalid combinations and set a global flag to true if one of the
# conditions is met

if week2_exists == True and week1_exists != True:
    somethings_wrong = True
if week3_exists == True and week1_exists != True:
    somethings_wrong = True
if week3_exists == True and week2_exists != True:
    somethings_wrong = True
if week4_exists == True and week1_exists != True:
    somethings_wrong = True
if week4_exists == True and week2_exists != True:
    somethings_wrong = True
if week4_exists == True and week3_exists != True:
    somethings_wrong = True

# finally, check if week4 exists and is full
# (sections checked =Num I added + 1, so a full set (ie 1 week) has 8 sections
if week4_exists == True:
    if num_sections_wk4 >= 8:
        clean_up_your_mess = True
        somethings_wrong = True

# if something has not gone wrong so far, set active document

document_set = False
set_to_doc1 = False
set_to_doc2 = False
set_to_doc3 = False
set_to_doc4 = False

if somethings_wrong == False:
    if week4_exists == True:
        document = document4
        document_set = True
        set_to_doc4 = True
    elif week3_exists == True and document_set == False:
        document = document3
        document_set = True
        set_to_doc3 = True
    elif week2_exists == True and document_set == False:
        document = document2
        document_set = True
        set_to_doc2 = True
    elif week1_exists == True and document_set == False:
        document = document1
        document_set = True
        set_to_doc1 = True
    else:
        document = Document()

# if something has not gone wrong so far, see if the document needs incremented
# and increment active document, creating empty next week

if somethings_wrong == False:
    if set_to_doc3 == True:
        if num_sections_wk3 >= 8:
            document4 = Document()
            document = document4
            set_to_doc4 = True
            # fix formatting
            section2 = document3.sections[-1]
            section2.orientation = WD_ORIENT.LANDSCAPE
            new_page_width, new_page_height = section2.page_height, section2.page_width
            section2.page_width = new_page_width
            section2.page_height = new_page_height
            # set margins
            section2.top_margin = Inches(0.5)
            section2.bottom_margin = Inches(0.5)
            section2.left_margin = Inches(0.5)
            section2.right_margin = Inches(0.5)
            document3.save("Week3.docx")

    if set_to_doc2 == True:
        if num_sections_wk2 >= 8:
            document3 = Document()
            document = document3
            set_to_doc3 = True
            # fix formatting
            section2 = document2.sections[-1]
            section2.orientation = WD_ORIENT.LANDSCAPE
            new_page_width, new_page_height = section2.page_height, section2.page_width
            section2.page_width = new_page_width
            section2.page_height = new_page_height
            # set margins
            section2.top_margin = Inches(0.5)
            section2.bottom_margin = Inches(0.5)
            section2.left_margin = Inches(0.5)
            section2.right_margin = Inches(0.5)
            document2.save("Week2.docx")

    if set_to_doc1 == True:
        if num_sections_wk1 >= 8:
            document2 = Document()
            document = document2
            set_to_doc2 = True
            # fix formatting
            section2 = document1.sections[-1]
            section2.orientation = WD_ORIENT.LANDSCAPE
            new_page_width, new_page_height = section2.page_height, section2.page_width
            section2.page_width = new_page_width
            section2.page_height = new_page_height
            # set margins
            section2.top_margin = Inches(0.5)
            section2.bottom_margin = Inches(0.5)
            section2.left_margin = Inches(0.5)
            section2.right_margin = Inches(0.5)
            document1.save("Week1.docx")

# actually start writing to the document

if somethings_wrong == False:
    # header paragraph
    header_paragraph = document.add_paragraph()
    report_date_display = header_paragraph.add_run("Report For: " + report_day)
    header_font = report_date_display.font
    header_font.name = "Times New Roman"
    header_font.size = Pt(22)

    # create invisible table for reports
    new_table = document.add_table(rows=2, cols=2)

    # add pie chart report
    pie_cell = new_table.cell(0, 0)
    pie_paragraph = pie_cell.paragraphs[0]
    pie_run = pie_paragraph.add_run()
    pie_run.add_picture("pie_report.png", width=Inches(6), height=Inches(3.76))

    # merge day report cells
    left_merged_cell = new_table.cell(1, 0)
    left_merged_cell.merge(new_table.cell(1, 1))

    # add line chart day report
    left_merged_cell_paragraph = left_merged_cell.paragraphs[0]
    day_run = left_merged_cell_paragraph.add_run()
    day_run.add_picture("day_report.png", width=Inches(10), height=Inches(2.32))

    document.add_paragraph()
    document.add_section(start_type=0)

    # new solution for page formatting
    section = document.sections[-2]
    section.orientation = WD_ORIENT.LANDSCAPE
    new_page_width, new_page_height = section.page_height, section.page_width
    section.page_width = new_page_width
    section.page_height = new_page_height

    # set margins
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)
    section.left_margin = Inches(0.5)
    section.right_margin = Inches(0.5)

    # end edge case: adding the last entry of the last day of the month:
    if week4_exists == True:
        num_sections_wk4 = len(document4.sections)
        if num_sections_wk4 >= 8:
            section2 = document4.sections[-1]
            section2.orientation = WD_ORIENT.LANDSCAPE
            new_page_width, new_page_height = section2.page_height, section2.page_width
            section2.page_width = new_page_width
            section2.page_height = new_page_height
            # set margins
            section2.top_margin = Inches(0.5)
            section2.bottom_margin = Inches(0.5)
            section2.left_margin = Inches(0.5)
            section2.right_margin = Inches(0.5)

    # document.add_picture("pie_report.png")
    # document.add_picture("day_report.png")
    if set_to_doc4:
        document.save("Week4.docx")
    elif set_to_doc3:
        document.save("Week3.docx")
    elif set_to_doc2:
        document.save("Week2.docx")
    elif set_to_doc1:
        document.save("Week1.docx")
    else:
        document.save('Week1.docx')
else:
    print("Something has gone horribly wrong with file storage, sort that out.")

if clean_up_your_mess:
    print("You are at the end of a reporting month. \nPlease move Week1 through",
          "Week4 to another folder for storage. \nThis program will not function",
          "until this action is completed.")

# saving process part 2: excel file for raw data
# open excel file and create workbook
new_sheet_name = report_day_converter(report_day)

try:
    wb = load_workbook("rawData.xlsx")
    ws = wb.create_sheet(new_sheet_name)
except:
    wb = Workbook()
    ws = wb.create_sheet(new_sheet_name)
    del wb["Sheet"]

# add labels to workbook
ws["A1"] = "Day Blocks"
ws["A2"] = "Start Time"
ws["B2"] = "End Time"
ws["C2"] = "Category"
ws["D2"] = "Duration"
ws["F1"] = "Categories by Duration"
ws["F2"] = "Category"
ws["G2"] = "Total Duration"

# add data to workbook
# add start times
start_time_counter = 3
for entry in chronological_activities:
    start_time = entry[0]
    converted_start_time = hour_and_minute_converter(start_time)
    ws.cell(row=start_time_counter, column=1, value=converted_start_time)
    start_time_counter += 1

# add end times
end_time_counter = 3
for entry in chronological_activities:
    end_time = entry[1]
    converted_end_time = hour_and_minute_converter(end_time)
    ws.cell(row=end_time_counter, column=2, value=converted_end_time)
    end_time_counter += 1

# add categories (day block table)
category_dayblock_counter = 3
for entry in chronological_activities:
    category = entry[2]
    ws.cell(row=category_dayblock_counter, column=3, value=category)
    category_dayblock_counter += 1

# add durations (day block table)
duration_counter = 3
for entry in chronological_activities:
    duration_minutes = entry[1] - entry[0] + 1
    duration_converted = hour_and_minute_converter(duration_minutes)
    ws.cell(row=duration_counter, column=4, value=duration_converted)
    duration_counter += 1

# add categories (total category duration table)
category_totalduration_counter = 3
for entry in table_categories:
    ws.cell(row=category_totalduration_counter, column=6, value=entry)
    category_totalduration_counter += 1

# add total durations
total_duration_counter = 3
for entry in table_durations:
    ws.cell(row=total_duration_counter, column=7, value=entry)
    total_duration_counter += 1

# save excel file
wb.save("rawData.xlsx")
